import requests
import logging
from PyPDF2 import PdfReader
from docx import Document
from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib import colors

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def load_content(file_path):
    """
    Load content from PDF, Word, or HTML file and return it as a string.
    """
    logger.info(f"Attempting to load content from: {file_path}")
    if file_path.endswith('.pdf'):
        return load_pdf(file_path)
    elif file_path.endswith('.docx'):
        return load_word(file_path)
    elif file_path.startswith('http'):
        return load_html(file_path)
    else:
        logger.error(f"Unsupported file format: {file_path}")
        raise ValueError("Unsupported file format")

def load_pdf(file_path):
    logger.info(f"Loading PDF file: {file_path}")
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            text = ""
            for i, page in enumerate(pdf_reader.pages):
                logger.debug(f"Extracting text from page {i+1}")
                text += page.extract_text()
        logger.info(f"Successfully loaded PDF file: {file_path}")
        return text
    except Exception as e:
        logger.error(f"Error loading PDF file {file_path}: {str(e)}")
        raise

def load_word(file_path):
    logger.info(f"Loading Word document: {file_path}")
    try:
        doc = Document(file_path)
        text = " ".join([paragraph.text for paragraph in doc.paragraphs])
        logger.info(f"Successfully loaded Word document: {file_path}")
        return text
    except Exception as e:
        logger.error(f"Error loading Word document {file_path}: {str(e)}")
        raise

def load_html(url):
    logger.info(f"Loading HTML content from URL: {url}")
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.content, 'html.parser')
        
        main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='content')
        
        if main_content:
            logger.info(f"Main content found and extracted from URL: {url}")
            return main_content.get_text(strip=True)
        else:
            logger.warning(f"No main content found, extracting all text from URL: {url}")
            return soup.get_text(strip=True)
    except Exception as e:
        logger.error(f"Error loading HTML content from {url}: {str(e)}")
        raise

def generate_comparison_output(original_text, translated_text, output_file, translation_method: str):
    """
    Generate a PDF or Word file with a side-by-side comparison of original and translated text.
    """
    logger.info(f"Generating comparison output: {output_file}")
    if output_file.endswith('.pdf'):
        output_file = output_file.replace('.pdf', '.docx')
    if output_file.endswith('.docx'):
        logger.info(f"Generating Word comparison: {output_file}")
        try:
            doc = Document()
            
            section = doc.sections[0]
            section.footer.add_paragraph(f"Translation automatically generated by Statistisches Amt Basel Stadt with {translation_method}")
            

            table = doc.add_table(rows=2, cols=2)
            
            table.cell(0, 0).text = "Original Text"
            table.cell(0, 1).text = "Translated Text"
            table.cell(1, 0).text = original_text
            table.cell(1, 1).text = translated_text

            doc.save(output_file)
            logger.info(f"Successfully generated Word comparison: {output_file}")
        except Exception as e:
            logger.error(f"Error generating Word comparison {output_file}: {str(e)}")
            raise
    else:
        raise ValueError("Unsuported file format. Only word output files are supported.")
